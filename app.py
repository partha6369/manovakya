import os
import re
import html
import fitz  # PyMuPDF
import gradio as gr
import google.generativeai as genai
from google.generativeai.types import content_types
from docx import Document

import nltk
from nltk.tokenize import sent_tokenize

nltk.download('punkt', quiet=True)
nltk.download('punkt_tab', quiet=True)

# === Load Gemini API Key ===
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", None)
if not GEMINI_API_KEY:
    raise ValueError("⚠️ Set GEMINI_API_KEY in environment or Hugging Face Secrets.")

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel(model_name="models/gemini-1.5-flash")

# === Helper Functions ===

def gemini_generate(prompt: str) -> str:
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"⚠️ Gemini API Error: {e}"

def generate_keywords(text: str, num_keywords: int) -> str:
    prompt = (
        "Analyse the following text. "
        f"Return {num_keywords} keywords or key phrases for the text, separated by a comma.\n\n"
        f"{text}"
    )
    return gemini_generate(prompt)

def generate_abstract(text: str, num_words: int) -> str:
    prompt = (
        "Analyse the following text. "
        f"Return an abstract for the text in approximately {num_words} words.\n\n"
        f"{text}"
    )
    return gemini_generate(prompt)

def generate_abstract_and_keywords(text: str, num_words: int, num_keywords: int) -> tuple[str, str]:
    abstract = generate_abstract(text, num_words)
    keywords = generate_keywords(text, num_keywords)

    return abstract, keywords
    
def analyze_text(text: str) -> str:
    prompt = (
        "Analyse the following text. "
        "Return sentiment (positive, negative, or neutral), topics, word count, and readability score.\n\n"
        f"{text}"
    )
    return gemini_generate(prompt)

def parse_analysis_output(response: str):
    sentiment = ""
    topics = ""
    word_count = ""
    readability = ""

    try:
        sentiment_match = re.search(r"\*\*Sentiment:\*\*\s*(.*)", response)
        topics_match = re.search(r"\*\*Topics:\*\*\s*(.*)", response)
        wordcount_match = re.search(r"\*\*Word Count:\*\*\s*(.*)", response)
        readability_match = re.search(r"\*\*Readability Score:\*\*\s*(.*)", response, re.DOTALL)

        sentiment = sentiment_match.group(1).strip() if sentiment_match else ""
        topics = topics_match.group(1).strip() if topics_match else ""
        word_count = wordcount_match.group(1).strip() if wordcount_match else ""
        readability = readability_match.group(1).strip() if readability_match else response.strip()
    except Exception as e:
        print(f"⚠️ Parsing error: {e}")
        readability = response.strip()

    return sentiment, topics, word_count, readability
    
def analyze_text_split_output(text):
    response = analyze_text(text)
    return parse_analysis_output(response)

def summarize_text(text: str) -> str:
    # Tokenise the input text into sentences
    sentences = sent_tokenize(text)
    total_sentences = len(sentences)

    # Compute 10% of the total, clamp between 2 and 5 sentences
    num_summary_sentences = max(2, int(total_sentences * 0.10))
    num_summary_sentences = min(num_summary_sentences, 5)

    prompt = (
        f"Summarise the following text in approximately {num_summary_sentences} sentences:\n\n{text}"
    )
    return gemini_generate(prompt)

def process_document(file):
    yield "⏳ *Processing document...*"

    if file is None:
        yield "❗ **No file uploaded.**"
        return

    ext = os.path.splitext(file.name)[-1].lower()
    try:
        if ext == ".pdf":
            try:
                uploaded_file = genai.upload_file(file.name, mime_type="application/pdf")
                response = model.generate_content(
                    ["Summarise this document:", uploaded_file]
                )
                text = response.text.strip()
            except Exception as e:
                yield f"⚠️ **Could not summarise PDF:** {e}"
                return

        elif ext == ".txt":
            try:
                # Try UTF-8 first
                with open(file.name, "r", encoding="utf-8") as f:
                    input_text = f.read()
            except UnicodeDecodeError:
                try:
                    # Fallback encoding
                    with open(file.name, "r", encoding="ISO-8859-1") as f:
                        input_text = f.read()
                except Exception as e:
                    yield f"⚠️ **Could not read TXT:** {e}"
                    return

            text = summarize_text(input_text)
            text = html.escape(text)

        elif ext == ".docx":
            try:
                doc = Document(file.name)
                input_text = "\n".join([para.text for para in doc.paragraphs])
                text = summarize_text(input_text)
                text = html.escape(text)
            except Exception as e:
                yield f"⚠️ **Could not process DOCX:** {e}"
                return

        else:
            yield "⚠️ **Only PDF, DOCX, and TXT files are supported.**"
            return

    except Exception as e:
        yield f"⚠️ **Unexpected error during processing:** {e}"
        return

    # Wrap the output in a scrollable div
    yield text
    
def process_file_for_abstract_and_keywords(file, num_words, num_keywords):
    yield "⏳ *Processing document...*", ""

    if file is None:
        yield "❗ **No file uploaded.**", ""
        return

    ext = os.path.splitext(file.name)[-1].lower()
    try:
        if ext == ".pdf":
            try:
                with fitz.open(file.name) as doc:
                    input_text = ""
                    for page in doc:
                        input_text += page.get_text()
        
                abstract, keywords = generate_abstract_and_keywords(input_text, num_words, num_keywords)
            except Exception as e:
                yield f"⚠️ **Could not generate abstract and extract keywords from PDF:** {e}", ""
                return

        elif ext == ".txt":
            try:
                with open(file.name, "r", encoding="utf-8") as f:
                    input_text = f.read()
            except UnicodeDecodeError:
                try:
                    with open(file.name, "r", encoding="ISO-8859-1") as f:
                        input_text = f.read()
                except Exception as e:
                    yield f"⚠️ **Could not read TXT:** {e}", ""
                    return

            abstract, keywords = generate_abstract_and_keywords(input_text, num_words, num_keywords)

        elif ext == ".docx":
            try:
                doc = Document(file.name)
                input_text = "\n".join([para.text for para in doc.paragraphs])
                abstract, keywords = generate_abstract_and_keywords(input_text, num_words, num_keywords)
            except Exception as e:
                yield f"⚠️ **Could not process DOCX:** {e}", ""
                return

        else:
            yield f"⚠️ **Unsupported file type `{ext}`.**", ""
            return

    except Exception as e:
        yield f"⚠️ **Unexpected error during processing:** {e}", ""
        return

    # Final result: two outputs (abstract and keywords)
    yield abstract, keywords
    
# ===== UI Title and Description =====
APP_TITLE = "🧘‍♂️ ManoVākya (मनॊवाक्य): Sentiments & Summaries"
APP_DESCRIPTION = (
    "ManoVākya (मनॊवाक्य) is your intelligent companion for understanding and interpreting language with clarity and insight. "
    "Whether you're expressing personal reflections, composing professional communication, or working with extensive documents, ManoVākya empowers you to uncover the tone, structure, sentiment, and essence of your words. "
    "With advanced AI capabilities, it can analyse emotional undertones, summarise lengthy texts into digestible insights, and help you discover key themes and readability in your writing. "
    "Designed for students, professionals, researchers, and reflective individuals alike, ManoVākya brings precision to journaling, content creation, documentation, and mindful communication. "
    "By combining linguistic intelligence with ethical design, ManoVākya bridges emotion and analysis—helping you not just write better, but think deeper."
)

# ===== Gradio Interfaces =====
analyze_interface = gr.Interface(
    fn=analyze_text_split_output,
    inputs=gr.Textbox(lines=5, max_lines=20, label="🗣 Enter text for sentiment & topic analysis"),
    outputs=[
        gr.Textbox(label="📊 Sentiment", lines = 1, max_lines=4),
        gr.Markdown(label="🧩 Topics", show_copy_button=True),
        gr.Textbox(label="🔢 Word Count", lines = 1),
        gr.Markdown(label="📘 Readability Score", show_copy_button=True)
    ],
    title="Gemini Sentiment & Topic Analyzer",
    flagging_mode="never",
    live=False
)

summarize_interface = gr.Interface(
    fn=summarize_text,
    inputs=gr.Textbox(lines=5, max_lines=20, label="✍️ Enter text to summarise"),
    outputs=gr.Textbox(label="📝 Summary Result", lines = 4, max_lines=20, show_copy_button=True),
    title="Gemini Text Summariser",
    flagging_mode="never",
    live=False
)

with gr.Blocks(title="Gemini Document Summariser") as doc_interface:
    gr.Markdown("## 📁 Gemini Document Summariser")

    file_input = gr.File(label="Upload PDF or DOCX or TXT File",
                         file_count = 'single',
                         file_types = ['.docx', '.txt', '.pdf']
                        )
    with gr.Row():
        submit_button = gr.Button("Submit")
        clear_button = gr.Button("Clear")
    output_box = gr.Markdown(label="📝 Summary", show_copy_button=True)

    submit_button.click(
        fn=process_document,
        inputs=file_input,
        outputs=output_box
    )

    clear_button.click(
        fn=lambda: (None, " "),  # Reset file and text
        inputs=[],
        outputs=[file_input, output_box]
    )

with gr.Blocks(title="Research Companion") as research_interface:
    gr.Markdown("## 🔍 Research Companion")
    
    file_input_r = gr.File(label="Upload PDF or DOCX or TXT File",
                           file_count = 'single',
                           file_types = ['.docx', '.txt', '.pdf']
                           )
    with gr.Row():
        txt_num_words_r = gr.Number(label="Number of Words in Abstract", value=300, precision=0)
        txt_num_keywords_r = gr.Number(label="Number of Keywords", value=5, precision=0)
    with gr.Row():
        submit_button_r = gr.Button("Generate Abstract and Keywords")
        clear_button_r = gr.Button("Clear")
    abstract_box_r = gr.Markdown(label="🧾 Abstract", show_copy_button=True)
    keywords_box_r = gr.Markdown(label="🔑 Keywords", show_copy_button=True)

    submit_button_r.click(
        fn=process_file_for_abstract_and_keywords,
        inputs=[file_input_r,
                txt_num_words_r,
                txt_num_keywords_r
               ],
        outputs=[abstract_box_r,
                 keywords_box_r
                ]
    )

    clear_button_r.click(
        fn=lambda: (None, " ", " "),  # Reset file, abstract, and keywords
        inputs=[],
        outputs=[file_input_r, abstract_box_r, keywords_box_r]
    )
    
# Try to load the PayPal URL from the environment; if missing, use a placeholder
paypal_url = os.getenv("PAYPAL_URL", "https://www.paypal.com/donate/dummy-link")

# ===== Launch Tabbed UI =====
with gr.Blocks(title=APP_TITLE) as the_application:
    gr.HTML(f"""
        <div style='text-align: center; margin-bottom: 20px;'>
            <p style='font-size: 40px; font-weight: bold;'>{APP_TITLE}</p>
            <p style='font-size: 16px; line-height: 1.6;'>{APP_DESCRIPTION}</p>
        </div>
    """)

    with gr.Tabs():
        with gr.TabItem("🗣 Analyse Text"):
            analyze_interface.render()
        with gr.TabItem("✍️ Summarise Text"):
            summarize_interface.render()
        with gr.TabItem("📁 Summarise Document"):
            doc_interface.render()
        with gr.TabItem("🔍 Research Companion"):
            research_interface.render()

    with gr.Row():
        gr.HTML(f"""
        <a href="{paypal_url}" target="_blank">
            <button style="background-color:#0070ba;color:white;border:none;padding:10px 20px;
            font-size:16px;border-radius:5px;cursor:pointer;margin-top:10px;">
                ❤️ Support Research via PayPal
            </button>
        </a>
        """)

    gr.Image(value="App Image.jpg", show_label=False, container=False)

# Determine if running on Hugging Face Spaces
on_spaces = os.environ.get("SPACE_ID") is not None

the_application.launch(share=not on_spaces)